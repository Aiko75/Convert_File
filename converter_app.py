from __future__ import annotations

import csv
import html
import io
import json
import os
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple
import xml.etree.ElementTree as ET

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText


@dataclass(frozen=True)
class ConversionJob:
    source_ext: str
    target_ext: str
    label: str
    action: Callable[[Path, Path], None]


class TextConverterApp:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title("Text Converter Studio")
        self.root.geometry("980x680")

        self.selected_source_files: List[Path] = []

        if getattr(sys, "frozen", False):
            app_base_dir = Path(sys.executable).resolve().parent
        else:
            app_base_dir = Path(__file__).resolve().parent

        self.workspace_dir = (app_base_dir / "converted_file" / "workspace").resolve()
        self.input_dir = self.workspace_dir / "input"
        self.output_dir = self.workspace_dir / "output"
        self.input_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.source_file_var = tk.StringVar()
        self.output_dir_var = tk.StringVar(value=str(self.output_dir))
        self.target_ext_var = tk.StringVar()

        self.jobs = self._build_jobs()
        self.supported_source_exts = sorted({job.source_ext for job in self.jobs})

        self._build_ui()

    def _build_ui(self) -> None:
        container = ttk.Frame(self.root, padding=12)
        container.pack(fill=tk.BOTH, expand=True)

        ttk.Label(container, text="File nguồn (có thể chọn nhiều file):").grid(row=0, column=0, sticky="w")
        ttk.Entry(container, textvariable=self.source_file_var, width=100).grid(
            row=1, column=0, sticky="we", padx=(0, 8)
        )
        ttk.Button(container, text="Chọn file...", command=self.pick_source_files).grid(row=1, column=1)

        ttk.Label(container, text="Thư mục output:").grid(row=2, column=0, sticky="w", pady=(10, 0))
        ttk.Entry(container, textvariable=self.output_dir_var, width=100).grid(
            row=3, column=0, sticky="we", padx=(0, 8)
        )
        ttk.Button(container, text="Chọn thư mục...", command=self.pick_output_dir).grid(row=3, column=1)

        ttk.Label(container, text="Định dạng đích:").grid(row=4, column=0, sticky="w", pady=(10, 0))
        self.target_combo = ttk.Combobox(
            container,
            textvariable=self.target_ext_var,
            state="readonly",
            width=30,
            values=[],
        )
        self.target_combo.grid(row=5, column=0, sticky="w")

        button_frame = ttk.Frame(container)
        button_frame.grid(row=6, column=0, sticky="w", pady=(10, 10))
        ttk.Button(button_frame, text="Convert", command=self.convert).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="Mở output", command=self.open_output_dir).pack(side=tk.LEFT, padx=(8, 0))

        ttk.Label(container, text="Các định dạng nguồn hỗ trợ:").grid(row=7, column=0, sticky="w")
        ttk.Label(
            container,
            text=", ".join(ext.upper().replace(".", "") for ext in self.supported_source_exts),
        ).grid(row=8, column=0, sticky="w", pady=(0, 6))

        ttk.Label(container, text="Log:").grid(row=9, column=0, sticky="w")
        self.log_text = ScrolledText(container, wrap=tk.WORD, height=20)
        self.log_text.grid(row=10, column=0, columnspan=2, sticky="nsew")

        container.columnconfigure(0, weight=1)
        container.rowconfigure(10, weight=1)

        self.log(f"Sẵn sàng. Folder làm việc: {self.workspace_dir}")
        self.log("App tập trung chuyển đổi văn bản và có hỗ trợ PDF <-> DOCX.")

    def _build_jobs(self) -> List[ConversionJob]:
        return [
            ConversionJob(".txt", ".md", "TXT -> MD", self.txt_to_md),
            ConversionJob(".md", ".txt", "MD -> TXT", self.md_to_txt),
            ConversionJob(".md", ".html", "MD -> HTML", self.md_to_html),
            ConversionJob(".html", ".md", "HTML -> MD", self.html_to_md),
            ConversionJob(".txt", ".html", "TXT -> HTML", self.txt_to_html),
            ConversionJob(".html", ".txt", "HTML -> TXT", self.html_to_txt),
            ConversionJob(".txt", ".docx", "TXT -> DOCX", self.txt_to_docx),
            ConversionJob(".docx", ".txt", "DOCX -> TXT", self.docx_to_txt),
            ConversionJob(".md", ".docx", "MD -> DOCX", self.md_to_docx),
            ConversionJob(".docx", ".md", "DOCX -> MD", self.docx_to_md),
            ConversionJob(".txt", ".pdf", "TXT -> PDF", self.txt_to_pdf),
            ConversionJob(".pdf", ".txt", "PDF -> TXT", self.pdf_to_txt),
            ConversionJob(".pdf", ".docx", "PDF -> DOCX", self.pdf_to_docx),
            ConversionJob(".docx", ".pdf", "DOCX -> PDF", self.docx_to_pdf),
            ConversionJob(".html", ".pdf", "HTML -> PDF", self.html_to_pdf),
            ConversionJob(".csv", ".json", "CSV -> JSON", self.csv_to_json),
            ConversionJob(".json", ".csv", "JSON -> CSV", self.json_to_csv),
            ConversionJob(".csv", ".txt", "CSV -> TXT", self.csv_to_txt),
            ConversionJob(".json", ".txt", "JSON -> TXT", self.json_to_txt),
            ConversionJob(".txt", ".json", "TXT -> JSON", self.txt_to_json),
        ]

    def log(self, message: str) -> None:
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def pick_source_files(self) -> None:
        filetypes = [
            (
                "Supported files",
                "*.txt *.md *.html *.docx *.pdf *.csv *.json",
            ),
            ("All files", "*.*"),
        ]
        selected = filedialog.askopenfilenames(
            title="Chọn một hoặc nhiều file nguồn",
            initialdir=str(self.input_dir),
            filetypes=filetypes,
        )
        if not selected:
            return

        self.selected_source_files = [Path(item) for item in selected]
        display_names = [item.name for item in self.selected_source_files]
        self.source_file_var.set("; ".join(str(item) for item in self.selected_source_files))

        source_exts = {item.suffix.lower() for item in self.selected_source_files}
        target_ext_sets = [
            {job.target_ext for job in self.jobs if job.source_ext == ext}
            for ext in source_exts
        ]
        target_exts = sorted(set.intersection(*target_ext_sets)) if target_ext_sets else []

        self.target_combo["values"] = target_exts
        if target_exts:
            self.target_ext_var.set(target_exts[0])
            self.log(
                "Đã nạp "
                f"{len(self.selected_source_files)} file ({', '.join(display_names[:4])}"
                f"{'...' if len(display_names) > 4 else ''}). "
                f"Có thể convert sang: {', '.join(target_exts)}"
            )
        else:
            self.target_ext_var.set("")
            self.log(
                "Các file đã chọn không có định dạng đích chung để convert hàng loạt. "
                "Hãy chọn nhóm file cùng kiểu."
            )

    def pick_output_dir(self) -> None:
        selected = filedialog.askdirectory(
            title="Chọn thư mục output",
            initialdir=self.output_dir_var.get() or str(self.output_dir),
        )
        if selected:
            self.output_dir_var.set(selected)

    def open_output_dir(self) -> None:
        output = Path(self.output_dir_var.get().strip())
        output.mkdir(parents=True, exist_ok=True)
        os.startfile(str(output))

    def convert(self) -> None:
        source_paths = self.selected_source_files
        target_ext = self.target_ext_var.get().strip().lower()
        output_dir = Path(self.output_dir_var.get().strip())

        if not source_paths:
            messagebox.showerror("Lỗi", "Bạn chưa chọn file nguồn.")
            return
        if not target_ext:
            messagebox.showerror("Lỗi", "Bạn chưa chọn định dạng đích.")
            return

        output_dir.mkdir(parents=True, exist_ok=True)
        success_count = 0
        failed_count = 0

        for source_path in source_paths:
            if not source_path.exists() or not source_path.is_file():
                failed_count += 1
                self.log(f"[LỖI] File không tồn tại: {source_path}")
                continue

            job = self.find_job(source_path.suffix.lower(), target_ext)
            if job is None:
                failed_count += 1
                self.log(
                    f"[BỎ QUA] {source_path.name}: không có converter "
                    f"{source_path.suffix.lower()} -> {target_ext}"
                )
                continue

            target_path = output_dir / f"{source_path.stem}{target_ext}"
            try:
                self.log(f"Bắt đầu: {job.label} | {source_path.name}")
                job.action(source_path, target_path)
                success_count += 1
                self.log(f"Hoàn tất: {target_path}")
            except Exception as error:
                failed_count += 1
                self.log(f"[LỖI] {source_path.name}: {error}")

        if failed_count == 0 and success_count > 0:
            messagebox.showinfo(
                "Thành công",
                f"Đã convert {success_count} file.\nOutput: {output_dir}",
            )
            return

        if success_count == 0:
            messagebox.showerror(
                "Lỗi convert",
                "Không convert được file nào. Xem log để biết chi tiết.",
            )
            return

        messagebox.showwarning(
            "Hoàn tất có lỗi",
            f"Convert thành công {success_count} file, lỗi {failed_count} file.\n"
            f"Output: {output_dir}",
        )

    def find_job(self, source_ext: str, target_ext: str) -> Optional[ConversionJob]:
        for job in self.jobs:
            if job.source_ext == source_ext and job.target_ext == target_ext:
                return job
        return None

    @staticmethod
    def read_text(path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def write_text(path: Path, text: str) -> None:
        path.write_text(text, encoding="utf-8")

    def txt_to_md(self, source: Path, target: Path) -> None:
        self.write_text(target, self.read_text(source))

    def md_to_txt(self, source: Path, target: Path) -> None:
        html_content = self._markdown_to_html(self.read_text(source))
        text = self._html_to_text(html_content)
        self.write_text(target, text)

    def md_to_html(self, source: Path, target: Path) -> None:
        html_content = self._markdown_to_html(self.read_text(source))
        self.write_text(target, html_content)

    def html_to_md(self, source: Path, target: Path) -> None:
        markdownify = self._import_markdownify()
        content = self.read_text(source)
        markdown_text = markdownify.markdownify(content, heading_style="ATX")
        self.write_text(target, markdown_text)

    def txt_to_html(self, source: Path, target: Path) -> None:
        text = self.read_text(source)
        escaped = html.escape(text)
        html_content = f"<html><body><pre>{escaped}</pre></body></html>"
        self.write_text(target, html_content)

    def html_to_txt(self, source: Path, target: Path) -> None:
        self.write_text(target, self._html_to_text(self.read_text(source)))

    def txt_to_docx(self, source: Path, target: Path) -> None:
        docx = self._import_docx()
        document = docx.Document()
        for line in self.read_text(source).splitlines():
            document.add_paragraph(line)
        document.save(target)

    def docx_to_txt(self, source: Path, target: Path) -> None:
        try:
            docx = self._import_docx()
            document = docx.Document(source)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            lines = self._extract_docx_text_without_python_docx(source)
            text = "\n".join(lines)
        self.write_text(target, text)

    def md_to_docx(self, source: Path, target: Path) -> None:
        docx = self._import_docx()
        markdown_text = self.read_text(source)
        document = docx.Document()
        for line in markdown_text.splitlines():
            stripped = line.strip()
            if stripped.startswith("### "):
                document.add_heading(stripped[4:].strip(), level=3)
            elif stripped.startswith("## "):
                document.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith("# "):
                document.add_heading(stripped[2:].strip(), level=1)
            else:
                document.add_paragraph(line)
        document.save(target)

    def docx_to_md(self, source: Path, target: Path) -> None:
        lines: List[str] = []
        try:
            docx = self._import_docx()
            document = docx.Document(source)
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    lines.append("")
                    continue
                style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
                if "heading 1" in style_name:
                    lines.append(f"# {text}")
                elif "heading 2" in style_name:
                    lines.append(f"## {text}")
                elif "heading 3" in style_name:
                    lines.append(f"### {text}")
                else:
                    lines.append(text)
        except Exception:
            lines = self._extract_docx_text_without_python_docx(source)
        self.write_text(target, "\n".join(lines))

    def txt_to_pdf(self, source: Path, target: Path) -> None:
        self._render_text_to_pdf(self.read_text(source), target)

    def pdf_to_txt(self, source: Path, target: Path) -> None:
        pypdf = self._import_pypdf()
        reader = pypdf.PdfReader(str(source))
        lines: List[str] = []
        for page in reader.pages:
            lines.append(page.extract_text() or "")
        self.write_text(target, "\n\n".join(lines).strip())

    def pdf_to_docx(self, source: Path, target: Path) -> None:
        pdf2docx = self._import_pdf2docx()
        converter = pdf2docx.Converter(str(source))
        try:
            converter.convert(str(target), start=0, end=None)
        finally:
            converter.close()

    def docx_to_pdf(self, source: Path, target: Path) -> None:
        try:
            docx2pdf = self._import_docx2pdf()
            original_stdout = sys.stdout
            original_stderr = sys.stderr
            restore_stdout = False
            restore_stderr = False

            try:
                if sys.stdout is None:
                    sys.stdout = io.StringIO()
                    restore_stdout = True
                if sys.stderr is None:
                    sys.stderr = io.StringIO()
                    restore_stderr = True

                docx2pdf.convert(str(source), str(target))
            finally:
                if restore_stdout:
                    sys.stdout = original_stdout
                if restore_stderr:
                    sys.stderr = original_stderr
            return
        except Exception as error:
            self.log(f"[INFO] DOCX -> PDF dùng docx2pdf thất bại, chuyển sang fallback: {error}")

        try:
            self._render_docx_with_format_to_pdf(source, target)
            return
        except Exception as error:
            self.log(f"[INFO] Fallback giữ format DOCX thất bại, dùng fallback text: {error}")

        try:
            docx = self._import_docx()
            document = docx.Document(source)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            lines = self._extract_docx_text_without_python_docx(source)
            text = "\n".join(lines)
        self._render_text_to_pdf(text, target)

    def html_to_pdf(self, source: Path, target: Path) -> None:
        text = self._html_to_text(self.read_text(source))
        self._render_text_to_pdf(text, target)

    def csv_to_json(self, source: Path, target: Path) -> None:
        with source.open("r", encoding="utf-8", errors="ignore", newline="") as csv_file:
            reader = csv.DictReader(csv_file)
            rows = list(reader)
        target.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")

    def json_to_csv(self, source: Path, target: Path) -> None:
        data = json.loads(self.read_text(source))
        if isinstance(data, dict):
            data = [data]
        if not isinstance(data, list):
            raise ValueError("JSON phải là object hoặc array object để convert sang CSV")

        keys: List[str] = []
        for item in data:
            if isinstance(item, dict):
                for key in item.keys():
                    if key not in keys:
                        keys.append(key)

        if not keys:
            raise ValueError("Không tìm thấy trường dữ liệu để ghi CSV")

        with target.open("w", encoding="utf-8", newline="") as csv_file:
            writer = csv.DictWriter(csv_file, fieldnames=keys)
            writer.writeheader()
            for row in data:
                writer.writerow(row if isinstance(row, dict) else {})

    def csv_to_txt(self, source: Path, target: Path) -> None:
        with source.open("r", encoding="utf-8", errors="ignore", newline="") as csv_file:
            reader = csv.reader(csv_file)
            lines = [" | ".join(row) for row in reader]
        self.write_text(target, "\n".join(lines))

    def json_to_txt(self, source: Path, target: Path) -> None:
        data = json.loads(self.read_text(source))
        self.write_text(target, json.dumps(data, ensure_ascii=False, indent=2))

    def txt_to_json(self, source: Path, target: Path) -> None:
        text = self.read_text(source)
        payload = {
            "content": text,
            "line_count": len(text.splitlines()),
            "char_count": len(text),
        }
        target.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    def _render_text_to_pdf(self, text: str, target: Path) -> None:
        reportlab_canvas, pdfmetrics, _, ReportLabColor = self._import_reportlab_toolkit()
        regular_font, _ = self._register_unicode_pdf_fonts(pdfmetrics)
        pdf = reportlab_canvas.Canvas(str(target))
        page_width, page_height = pdf._pagesize
        left = 40
        right = 40
        top = page_height - 40
        bottom = 40
        font_size = 11
        line_height = 16
        max_width = page_width - left - right

        y = top
        pdf.setFillColor(ReportLabColor(0, 0, 0))
        pdf.setFont(regular_font, font_size)

        for raw_line in text.splitlines() or [""]:
            chunks = self._split_text_by_width(raw_line, regular_font, font_size, max_width, pdfmetrics)
            if not chunks:
                chunks = [""]

            for chunk in chunks:
                pdf.drawString(left, y, chunk)
                y -= line_height
                if y < bottom:
                    pdf.showPage()
                    y = top
                    pdf.setFillColor(ReportLabColor(0, 0, 0))
                    pdf.setFont(regular_font, font_size)
        pdf.save()

    def _render_docx_with_format_to_pdf(self, source: Path, target: Path) -> None:
        docx = self._import_docx()
        document = docx.Document(source)

        reportlab_canvas, pdfmetrics, _, ReportLabColor = self._import_reportlab_toolkit()
        regular_font, bold_font = self._register_unicode_pdf_fonts(pdfmetrics)

        pdf = reportlab_canvas.Canvas(str(target))
        page_width, page_height = pdf._pagesize
        left = 40
        right = 40
        top = page_height - 40
        bottom = 40
        font_size = 11
        line_height = 16
        max_width = page_width - left - right

        x = left
        y = top

        def ensure_page() -> None:
            nonlocal x, y
            if y >= bottom:
                return
            pdf.showPage()
            x = left
            y = top

        for paragraph in document.paragraphs:
            runs = list(paragraph.runs)
            if not runs:
                y -= line_height
                x = left
                ensure_page()
                continue

            for run in runs:
                run_text = run.text or ""
                if not run_text:
                    continue

                font_name = bold_font if bool(run.bold or run.font.bold) else regular_font
                text_rgb = (0, 0, 0)
                if run.font.color is not None and run.font.color.rgb is not None:
                    rgb = run.font.color.rgb
                    text_rgb = (rgb[0], rgb[1], rgb[2])

                highlight_rgb = self._docx_highlight_to_rgb(run.font.highlight_color)

                fragments = run_text.split("\n")
                for frag_index, fragment in enumerate(fragments):
                    tokens = re.split(r"(\s+)", fragment)
                    for token in tokens:
                        if token == "":
                            continue

                        token_width = pdfmetrics.stringWidth(token, font_name, font_size)
                        if x + token_width > left + max_width and x > left:
                            x = left
                            y -= line_height
                            ensure_page()

                        if highlight_rgb is not None and token.strip():
                            hl_color = ReportLabColor(
                                highlight_rgb[0] / 255.0,
                                highlight_rgb[1] / 255.0,
                                highlight_rgb[2] / 255.0,
                            )
                            pdf.setFillColor(hl_color)
                            pdf.rect(x, y - 2, token_width, line_height, fill=1, stroke=0)

                        txt_color = ReportLabColor(
                            text_rgb[0] / 255.0,
                            text_rgb[1] / 255.0,
                            text_rgb[2] / 255.0,
                        )
                        pdf.setFillColor(txt_color)
                        pdf.setFont(font_name, font_size)
                        pdf.drawString(x, y, token)
                        x += token_width

                    if frag_index < len(fragments) - 1:
                        x = left
                        y -= line_height
                        ensure_page()

            x = left
            y -= line_height
            ensure_page()

        pdf.save()

    def _markdown_to_html(self, markdown_text: str) -> str:
        markdown = self._import_markdown()
        return markdown.markdown(markdown_text)

    def _html_to_text(self, html_content: str) -> str:
        bs4 = self._import_bs4()
        soup = bs4.BeautifulSoup(html_content, "html.parser")
        text = soup.get_text("\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _extract_docx_text_without_python_docx(self, source: Path) -> List[str]:
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        lines: List[str] = []

        with zipfile.ZipFile(source, "r") as archive:
            if "word/document.xml" not in archive.namelist():
                raise RuntimeError("File DOCX không hợp lệ hoặc thiếu word/document.xml")

            xml_bytes = archive.read("word/document.xml")
            root = ET.fromstring(xml_bytes)

            for paragraph in root.findall(".//w:p", namespace):
                text_nodes = paragraph.findall(".//w:t", namespace)
                line = "".join(node.text or "" for node in text_nodes).strip()
                lines.append(line)

        if not lines:
            raise RuntimeError("Không trích xuất được nội dung từ file DOCX")

        return lines

    def _docx_highlight_to_rgb(self, highlight) -> Optional[Tuple[int, int, int]]:
        if highlight is None:
            return None

        name = getattr(highlight, "name", str(highlight)).upper()
        mapping: Dict[str, Tuple[int, int, int]] = {
            "YELLOW": (255, 255, 0),
            "BRIGHT_GREEN": (0, 255, 0),
            "TURQUOISE": (0, 255, 255),
            "PINK": (255, 192, 203),
            "BLUE": (173, 216, 230),
            "RED": (255, 0, 0),
            "DARK_BLUE": (0, 0, 139),
            "TEAL": (0, 128, 128),
            "GREEN": (0, 128, 0),
            "VIOLET": (238, 130, 238),
            "GRAY_25": (211, 211, 211),
            "GRAY_50": (169, 169, 169),
            "DARK_YELLOW": (204, 153, 0),
        }
        for key, value in mapping.items():
            if key in name:
                return value
        return (255, 255, 0)

    def _split_text_by_width(
        self,
        text: str,
        font_name: str,
        font_size: int,
        max_width: float,
        pdfmetrics,
    ) -> List[str]:
        if not text:
            return [""]

        tokens = re.split(r"(\s+)", text)
        lines: List[str] = []
        current = ""

        for token in tokens:
            if token == "":
                continue
            candidate = current + token
            width = pdfmetrics.stringWidth(candidate, font_name, font_size)
            if width <= max_width or not current:
                current = candidate
                continue
            lines.append(current.rstrip())
            current = token.lstrip()

        if current:
            lines.append(current.rstrip())

        return lines

    def _register_unicode_pdf_fonts(self, pdfmetrics) -> Tuple[str, str]:
        try:
            pdfmetrics.getFont("AppUnicode")
            pdfmetrics.getFont("AppUnicodeBold")
            return "AppUnicode", "AppUnicodeBold"
        except Exception:
            pass

        _, _, TTFont, _ = self._import_reportlab_toolkit()

        font_candidates = [
            ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf"),
            ("C:/Windows/Fonts/tahoma.ttf", "C:/Windows/Fonts/tahomabd.ttf"),
            ("C:/Windows/Fonts/times.ttf", "C:/Windows/Fonts/timesbd.ttf"),
        ]

        for regular_path, bold_path in font_candidates:
            if Path(regular_path).exists():
                pdfmetrics.registerFont(TTFont("AppUnicode", regular_path))
                if Path(bold_path).exists():
                    pdfmetrics.registerFont(TTFont("AppUnicodeBold", bold_path))
                else:
                    pdfmetrics.registerFont(TTFont("AppUnicodeBold", regular_path))
                return "AppUnicode", "AppUnicodeBold"

        return "Helvetica", "Helvetica-Bold"

    @staticmethod
    def _import_docx():
        try:
            import docx  # type: ignore
        except Exception as error:
            raise RuntimeError(
                "Không import được python-docx. Cài: pip install python-docx. "
                f"Chi tiết: {error}"
            ) from error
        return docx

    @staticmethod
    def _import_markdown():
        try:
            import markdown  # type: ignore
        except Exception as error:
            raise RuntimeError("Thiếu package markdown. Cài: pip install markdown") from error
        return markdown

    @staticmethod
    def _import_markdownify():
        try:
            import markdownify  # type: ignore
        except Exception as error:
            raise RuntimeError("Thiếu package markdownify. Cài: pip install markdownify") from error
        return markdownify

    @staticmethod
    def _import_bs4():
        try:
            import bs4  # type: ignore
        except Exception as error:
            raise RuntimeError("Thiếu package beautifulsoup4. Cài: pip install beautifulsoup4") from error
        return bs4

    @staticmethod
    def _import_pypdf():
        try:
            import pypdf  # type: ignore
        except Exception as error:
            raise RuntimeError("Thiếu package pypdf. Cài: pip install pypdf") from error
        return pypdf

    @staticmethod
    def _import_pdf2docx():
        try:
            import pdf2docx  # type: ignore
        except Exception as error:
            raise RuntimeError("Thiếu package pdf2docx. Cài: pip install pdf2docx") from error
        return pdf2docx

    @staticmethod
    def _import_docx2pdf():
        try:
            import docx2pdf  # type: ignore
        except Exception as error:
            raise RuntimeError("Thiếu package docx2pdf. Cài: pip install docx2pdf") from error
        return docx2pdf

    @staticmethod
    def _import_reportlab_toolkit():
        try:
            from reportlab.pdfgen import canvas  # type: ignore
            from reportlab.pdfbase import pdfmetrics  # type: ignore
            from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
            from reportlab.lib.colors import Color  # type: ignore
        except Exception as error:
            raise RuntimeError("Thiếu package reportlab. Cài: pip install reportlab") from error
        return canvas, pdfmetrics, TTFont, Color


def main() -> None:
    root = tk.Tk()
    app = TextConverterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
